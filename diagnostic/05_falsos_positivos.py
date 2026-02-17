"""
ETAPA 5 - Falsos Positivos: artigos excluidos pelo humano, porem
incluidos (maybe/include) pela IA.

Le o arquivo pareado gerado pela Etapa 1 (output/pareamento.xlsx),
filtra os falsos positivos e gera um relatorio em Word (.docx).

Uso:
    python 05_falsos_positivos.py
    python 05_falsos_positivos.py --input output/pareamento.xlsx
"""

import argparse
import sys
import datetime
import json
from pathlib import Path

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ------------------------------------------------------------------ paths --
SCRIPT_DIR  = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
OUTPUT_DIR  = PROJECT_DIR / "output"
DEFAULT_INPUT = OUTPUT_DIR / "pareamento.xlsx"


# ---------------------------------------------------------- normalisation --

def normalise_decision(s) -> str:
    """Padroniza variantes: included->include, excluded->exclude."""
    if pd.isna(s):
        return ""
    d = str(s).strip().lower()
    if d == "included":
        return "include"
    if d == "excluded":
        return "exclude"
    return d


# ======================================================= Word helpers ==

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=Pt(9), font_name="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)


def add_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_heading(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


# ========================================== Word document generation ==

def generate_word(df_fp, total_paired, output_dir: Path):
    """Generate Word document listing false positives."""

    doc = Document()

    # --- Document style defaults ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ================================================================
    #  TITLE
    # ================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("False Positives - Articles Excluded by Humans but Included by AI")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(f"Generated: {ts}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    n_fp = len(df_fp)

    # ================================================================
    #  TABLE 1 - Summary
    # ================================================================
    add_heading(doc, "Table 1. Summary")

    t1 = doc.add_table(rows=4, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t1)

    pct = f"{n_fp / total_paired * 100:.1f}%" if total_paired > 0 else "-"

    t1_data = [
        ("Characteristic", "Value"),
        ("Total paired articles", str(total_paired)),
        ("False positives (AI included, Human excluded)", str(n_fp)),
        ("Proportion of total", pct),
    ]
    for i, (left, right) in enumerate(t1_data):
        is_header = (i == 0)
        set_cell_text(t1.cell(i, 0), left, bold=is_header,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(t1.cell(i, 1), right, bold=is_header)
        if is_header:
            set_cell_shading(t1.cell(i, 0), "D9E2F3")
            set_cell_shading(t1.cell(i, 1), "D9E2F3")

    doc.add_paragraph()

    # ================================================================
    #  TABLE 2 - List of False Positive Articles
    # ================================================================
    add_heading(doc, "Table 2. False Positive Articles")

    note = doc.add_paragraph(
        "Articles that human reviewers excluded but the AI classified as "
        "maybe or include during title/abstract screening."
    )
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.italic = True

    # Determine columns to show
    has_id = "id" in df_fp.columns
    has_abstract = "abstract" in df_fp.columns
    has_rationale = "screening_reason" in df_fp.columns

    # Build header
    headers = ["#"]
    if has_id:
        headers.append("ID")
    headers.append("Title")
    headers.append("AI Decision")
    if has_rationale:
        headers.append("AI Rationale")

    n_cols = len(headers)
    t2 = doc.add_table(rows=1, cols=n_cols)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t2)

    for j, h in enumerate(headers):
        set_cell_text(t2.cell(0, j), h, bold=True)
        set_cell_shading(t2.cell(0, j), "D9E2F3")

    # Data rows
    for i, (_, row) in enumerate(df_fp.iterrows(), start=1):
        r = t2.add_row()
        col = 0

        # #
        set_cell_text(r.cells[col], str(i), align=WD_ALIGN_PARAGRAPH.CENTER)
        col += 1

        # ID
        if has_id:
            rid = str(row.get("id", ""))
            set_cell_text(r.cells[col], rid, align=WD_ALIGN_PARAGRAPH.CENTER)
            col += 1

        # Title
        title = str(row.get("title", ""))
        if len(title) > 200:
            title = title[:200] + "..."
        set_cell_text(r.cells[col], title, align=WD_ALIGN_PARAGRAPH.LEFT,
                      font_size=Pt(8))
        col += 1

        # AI Decision
        decision = str(row.get("screening_decision", ""))
        set_cell_text(r.cells[col], decision, align=WD_ALIGN_PARAGRAPH.CENTER)
        col += 1

        # AI Rationale
        if has_rationale:
            rationale = str(row.get("screening_reason", ""))
            if rationale == "nan":
                rationale = ""
            set_cell_text(r.cells[col], rationale,
                          align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(8))
            col += 1

    doc.add_paragraph()

    # ================================================================
    #  TABLE 3 - Articles with Abstract (for detailed review)
    # ================================================================
    if has_abstract:
        add_heading(doc, "Table 3. Detailed View (with Abstracts)")

        note2 = doc.add_paragraph(
            "Full title and abstract of each false positive article "
            "for detailed review."
        )
        note2.runs[0].font.size = Pt(8)
        note2.runs[0].font.italic = True

        for i, (_, row) in enumerate(df_fp.iterrows(), start=1):
            title = str(row.get("title", ""))
            abstract = str(row.get("abstract", ""))
            if abstract == "nan":
                abstract = "(no abstract)"
            decision = str(row.get("screening_decision", ""))
            rationale = str(row.get("screening_reason", ""))
            if rationale == "nan":
                rationale = ""

            # Article number + title
            p_title = doc.add_paragraph()
            run_num = p_title.add_run(f"{i}. ")
            run_num.bold = True
            run_num.font.size = Pt(10)
            run_num.font.name = "Times New Roman"
            run_t = p_title.add_run(title)
            run_t.bold = True
            run_t.font.size = Pt(10)
            run_t.font.name = "Times New Roman"

            # Decision + rationale
            p_dec = doc.add_paragraph()
            run_dec = p_dec.add_run(f"AI Decision: {decision}")
            run_dec.font.size = Pt(9)
            run_dec.font.name = "Times New Roman"
            run_dec.font.color.rgb = RGBColor(180, 0, 0)
            if rationale:
                run_rat = p_dec.add_run(f"  |  Rationale: {rationale}")
                run_rat.font.size = Pt(9)
                run_rat.font.name = "Times New Roman"
                run_rat.font.italic = True

            # Abstract
            p_abs = doc.add_paragraph()
            run_abs = p_abs.add_run(abstract)
            run_abs.font.size = Pt(8)
            run_abs.font.name = "Times New Roman"
            run_abs.font.color.rgb = RGBColor(80, 80, 80)
            p_abs.paragraph_format.space_after = Pt(10)

    # ================================================================
    #  SAVE
    # ================================================================
    ts_file = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = output_dir / f"falsos_positivos_{ts_file}.docx"
    doc.save(str(docx_path))
    return docx_path


# ========================================== console report ==

def print_console(df_fp, total_paired):
    n_fp = len(df_fp)
    pct = f"{n_fp / total_paired * 100:.1f}%" if total_paired > 0 else "-"

    print()
    print("=" * 62)
    print("  ETAPA 5 - FALSOS POSITIVOS")
    print("  Artigos excluidos pelo humano, incluidos pela IA")
    print("=" * 62)

    print(f"\n  Total pareados:          {total_paired}")
    print(f"  Falsos positivos (FP):   {n_fp}  ({pct})")

    if n_fp > 0:
        print(f"\n  ** Lista dos {n_fp} falsos positivos:")
        for i, (_, row) in enumerate(df_fp.iterrows(), start=1):
            title = str(row.get("title", ""))[:80]
            decision = str(row.get("screening_decision", ""))
            print(f"     {i:3d}. [{decision}] {title}")

    print("=" * 62)


# ------------------------------------------------------------------ main --

def main():
    parser = argparse.ArgumentParser(
        description=(
            "Etapa 5 - Falsos Positivos.\n\n"
            "Lista artigos que os humanos excluiram mas a IA incluiu.\n"
            "Le o arquivo pareado (output/pareamento.xlsx) gerado pela Etapa 1."
        ),
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument("--input", "-i", default=None,
                        help="Arquivo pareado XLSX/CSV (padrao: output/pareamento.xlsx).")
    args = parser.parse_args()

    input_path = Path(args.input) if args.input else DEFAULT_INPUT

    if not input_path.is_file():
        print(f"\n  ERRO: Arquivo pareado nao encontrado: {input_path}")
        print("  Rode primeiro: python 01_pareamento.py\n")
        sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load paired data
    ext = input_path.suffix.lower()
    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(str(input_path))
    else:
        df = pd.read_csv(str(input_path), encoding="utf-8-sig")

    # Normalise columns
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    if "screening_decision" not in df.columns or "decision_human" not in df.columns:
        print("  ERRO: Colunas 'screening_decision' e 'decision_human' nao encontradas.")
        print(f"  Colunas do arquivo: {list(df.columns)}")
        sys.exit(1)

    print(f"\n  Arquivo lido: {input_path.name}  ({len(df)} registros)")

    total_paired = len(df)

    # Normalise decisions
    df["screening_decision"] = df["screening_decision"].apply(normalise_decision)
    df["decision_human"] = df["decision_human"].apply(normalise_decision)

    # Filter false positives:
    #   Human = exclude  AND  AI = maybe or include
    fp_mask = (
        (df["decision_human"] == "exclude") &
        (df["screening_decision"].isin(["maybe", "include"]))
    )
    df_fp = df[fp_mask].copy().reset_index(drop=True)

    # Console
    print_console(df_fp, total_paired)

    # Word document
    docx_path = generate_word(df_fp, total_paired, OUTPUT_DIR)
    print(f"\n  Documento Word gerado: {docx_path.name}")
    print(f"  Pasta: {OUTPUT_DIR}")

    # Excel with false positives
    ts_file = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    xlsx_path = OUTPUT_DIR / f"falsos_positivos_{ts_file}.xlsx"
    df_fp.to_excel(str(xlsx_path), index=False)
    print(f"  Planilha Excel: {xlsx_path.name}")

    # JSON summary
    json_path = OUTPUT_DIR / f"falsos_positivos_{ts_file}.json"
    summary = {
        "total_paired": total_paired,
        "false_positives": len(df_fp),
        "proportion": round(len(df_fp) / total_paired, 4) if total_paired > 0 else 0,
        "articles": [
            {
                "title": str(row.get("title", "")),
                "ai_decision": str(row.get("screening_decision", "")),
                "ai_rationale": str(row.get("screening_reason", "")),
            }
            for _, row in df_fp.iterrows()
        ],
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(f"  Resumo JSON: {json_path.name}\n")


if __name__ == "__main__":
    main()

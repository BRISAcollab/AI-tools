"""
ETAPA 3 - Verificacao de artigos incluidos no texto completo (Fulltext Check).

Compara os artigos incluidos na revisao final (pos-fulltext) com a
decisao da IA na fase de TIAB. Avalia se a IA teria mantido esses
artigos para a fase seguinte (screening_decision = maybe/include).

Uso:
    1. Coloque na pasta input/:
       - O arquivo da IA (com screening_decision)
       - O arquivo de artigos incluidos na revisao final (fulltext)
    2. Rode:  python diagnostic/03_fulltext_check.py

    Modo manual:
        python diagnostic/03_fulltext_check.py --ai input/ai.xlsx --fulltext input/fulltext.xlsx
"""

import argparse
import sys
import os
import json
import datetime
from pathlib import Path

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ------------------------------------------------------------------ paths --
SCRIPT_DIR  = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
INPUT_DIR   = PROJECT_DIR / "input"
OUTPUT_DIR  = PROJECT_DIR / "output"


# ---------------------------------------------------------------- helpers --

def ensure_folders():
    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)


def load_file(path: str) -> pd.DataFrame:
    ext = Path(path).suffix.lower()
    if ext == ".csv":
        try:
            return pd.read_csv(path, encoding="utf-8")
        except UnicodeDecodeError:
            return pd.read_csv(path, encoding="latin-1")
    elif ext in (".xlsx", ".xls"):
        return pd.read_excel(path)
    else:
        raise ValueError(f"Formato nao suportado: {ext}. Use .csv, .xlsx ou .xls")


def normalise_columns(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    return df


def normalise_title(s) -> str:
    if pd.isna(s):
        return ""
    return str(s).strip().lower()


def normalise_decision(s) -> str:
    """Padroniza variantes: included->include, excluded->exclude."""
    if pd.isna(s):
        return ""
    d = str(s).strip().lower()
    if d == "included":
        return "include"
    if d == "excluded":
        return "exclude"
    return d


def auto_detect_files():
    """
    Detecta automaticamente o arquivo da IA e o arquivo fulltext em input/.
    - IA: contem 'screening_decision' e muitas linhas (>50)
    - Fulltext: poucos artigos, com 'title'
    """
    valid_ext = {".csv", ".xlsx", ".xls"}
    candidates = [f for f in INPUT_DIR.iterdir() if f.suffix.lower() in valid_ext]

    if len(candidates) < 2:
        print(f"\n  ERRO: Sao necessarios pelo menos 2 arquivos em '{INPUT_DIR}'.")
        sys.exit(1)

    ai_file = None
    fulltext_file = None

    scored = []
    for f in candidates:
        try:
            df = normalise_columns(load_file(str(f)))
            cols = set(df.columns)
            n = len(df)
        except Exception:
            continue

        has_screening = "screening_decision" in cols
        has_decision  = "decision" in cols
        has_id        = "id" in cols
        has_incl      = "inclusion_evaluation" in cols

        scored.append({
            "path": f,
            "n": n,
            "has_screening": has_screening,
            "has_decision": has_decision,
            "has_id": has_id,
            "has_incl": has_incl,
        })

    # AI file: has screening_decision + large number of records
    ai_candidates = [s for s in scored if s["has_screening"] and s["n"] > 50]
    ai_candidates.sort(key=lambda x: x["n"], reverse=True)
    if ai_candidates:
        ai_file = ai_candidates[0]["path"]

    # Fulltext file: smallest file that is NOT the AI file
    remaining = [s for s in scored if s["path"] != ai_file]
    # Prefer files with 'decision' column and fewer records
    ft_candidates = sorted(remaining, key=lambda x: x["n"])
    if ft_candidates:
        fulltext_file = ft_candidates[0]["path"]

    if ai_file is None or fulltext_file is None:
        print("  ERRO: Nao foi possivel detectar os arquivos automaticamente.")
        print("  Use: python diagnostic/03_fulltext_check.py --ai <arq_ia> --fulltext <arq_fulltext>")
        sys.exit(1)

    return str(ai_file), str(fulltext_file)


# --------------------------------------------------------------- analysis --

def run_fulltext_check(ai_path: str, fulltext_path: str, output_dir: Path):
    """Compara artigos incluidos no fulltext com a decisao da IA no TIAB."""

    ai_df = normalise_columns(load_file(ai_path))
    ft_df = normalise_columns(load_file(fulltext_path))

    # Detectar coluna de decisao da IA
    ai_decision_col = None
    for c in ("screening_decision", "screening", "decision_ai", "ai_decision"):
        if c in ai_df.columns:
            ai_decision_col = c
            break
    if ai_decision_col is None:
        raise KeyError(f"Coluna de decisao da IA nao encontrada. Colunas: {list(ai_df.columns)}")

    if "title" not in ai_df.columns:
        raise KeyError(f"Arquivo da IA sem coluna 'title'. Colunas: {list(ai_df.columns)}")
    if "title" not in ft_df.columns:
        raise KeyError(f"Arquivo fulltext sem coluna 'title'. Colunas: {list(ft_df.columns)}")

    # Normalizar titulos
    ai_df["_title_key"] = ai_df["title"].apply(normalise_title)
    ft_df["_title_key"] = ft_df["title"].apply(normalise_title)

    # Merge: procurar cada artigo fulltext dentro da base da IA
    merged = pd.merge(
        ft_df[["title", "_title_key"]],
        ai_df[["_title_key", ai_decision_col, "title"]],
        on="_title_key",
        how="left",
        suffixes=("_ft", "_ai"),
    )

    # Se houve match, usar a decisao da IA; senao, marcar como "nao encontrado"
    merged["ai_decision"] = merged[ai_decision_col].fillna("nao encontrado")
    merged["ai_decision_lower"] = merged["ai_decision"].apply(
        lambda x: normalise_decision(x) if x != "nao encontrado" else x
    )

    # Binarizar: include e maybe = "passaria" (positivo)
    merged["ai_passaria"] = merged["ai_decision_lower"].isin(["maybe", "include"])

    # Titulo para exibicao (preferir o do fulltext)
    merged["title_display"] = merged["title_ft"].fillna(merged["title_ai"])

    # ================================================================
    #  RESULTADOS
    # ================================================================
    total_fulltext  = len(ft_df)
    matched         = merged["ai_decision_lower"] != "nao encontrado"
    n_matched       = matched.sum()
    n_not_found     = total_fulltext - n_matched

    n_passaria      = merged["ai_passaria"].sum()
    n_perdidos      = n_matched - n_passaria  # matched mas IA excluiu

    lost_df = merged[matched & ~merged["ai_passaria"]].copy()
    found_df = merged[matched & merged["ai_passaria"]].copy()
    not_found_df = merged[~matched].copy()

    # Captura rate: dos artigos finais, quantos a IA teria mantido?
    capture_rate = n_passaria / n_matched if n_matched > 0 else float("nan")
    miss_rate    = n_perdidos / n_matched if n_matched > 0 else float("nan")

    # ================================================================
    #  CONSOLE
    # ================================================================
    print()
    print("=" * 62)
    print("  ETAPA 3 - FULLTEXT CHECK")
    print("  Artigos da revisao final vs decisao da IA no TIAB")
    print("=" * 62)

    print(f"\n  Artigos incluidos na revisao final:  {total_fulltext}")
    print(f"  Encontrados na base da IA:           {n_matched}")
    print(f"  Nao encontrados na base da IA:       {n_not_found}")

    if n_not_found > 0:
        print(f"\n  ** Artigos do fulltext NAO encontrados na IA:")
        for _, row in not_found_df.iterrows():
            print(f"     - {str(row['title_display'])[:90]}")

    print(f"\n  RESULTADO DA COMPARACAO")
    print("-" * 62)
    print(f"  IA manteria (maybe/include):  {n_passaria}/{n_matched}  ({capture_rate*100:.1f}%)")
    print(f"  IA perderia (exclude):        {n_perdidos}/{n_matched}  ({miss_rate*100:.1f}%)")

    if n_perdidos > 0:
        print(f"\n  ** Artigos PERDIDOS pela IA (excluidos no TIAB mas incluidos na revisao final):")
        for _, row in lost_df.iterrows():
            t = str(row["title_display"])[:90]
            d = row["ai_decision"]
            print(f"     - [{d}] {t}")

    if n_passaria == n_matched and n_matched > 0:
        print(f"\n  EXCELENTE: A IA teria mantido TODOS os {n_matched} artigos da revisao final!")

    print("=" * 62)

    # ================================================================
    #  GERAR TABELA DETALHADA (XLSX)
    # ================================================================
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    detail_df = merged[["title_display", "ai_decision", "ai_passaria"]].copy()
    detail_df.columns = ["title", "ai_tiab_decision", "ai_manteria"]
    detail_df["ai_manteria"] = detail_df["ai_manteria"].map({True: "Sim", False: "Nao"})
    detail_df.insert(0, "n", range(1, len(detail_df) + 1))

    xlsx_detail = output_dir / f"fulltext_check_{ts}.xlsx"
    detail_df.to_excel(xlsx_detail, index=False)

    # ================================================================
    #  GERAR DOCUMENTO WORD
    # ================================================================
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Fulltext Capture Analysis - AI TIAB Screening")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # --- Table 1: Summary ---
    h = doc.add_heading("Table 1. Fulltext Capture Summary", level=2)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)

    t1 = doc.add_table(rows=1, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(t1)

    for j, txt in enumerate(["Parameter", "Value"]):
        _set_cell(t1.cell(0, j), txt, bold=True)
        _shade(t1.cell(0, j), "D9E2F3")

    summary_rows = [
        ("Articles included in final review",      str(total_fulltext)),
        ("Found in AI TIAB screening",             str(n_matched)),
        ("Not found in AI database",               str(n_not_found)),
        ("AI would keep (maybe/include)",           f"{n_passaria} ({capture_rate*100:.1f}%)" if n_matched else "N/A"),
        ("AI would miss (exclude)",                 f"{n_perdidos} ({miss_rate*100:.1f}%)" if n_matched else "N/A"),
        ("Capture Rate",                            f"{capture_rate*100:.1f}%" if n_matched else "N/A"),
    ]
    for label, val in summary_rows:
        row = t1.add_row()
        _set_cell(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row.cells[1], val)

    doc.add_paragraph()

    # --- Table 2: Article-level detail ---
    h2 = doc.add_heading("Table 2. Article-Level Results", level=2)
    for r in h2.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)

    t2 = doc.add_table(rows=1, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(t2)

    for j, txt in enumerate(["#", "Article Title", "AI TIAB Decision"]):
        _set_cell(t2.cell(0, j), txt, bold=True)
        _shade(t2.cell(0, j), "D9E2F3")

    for i, (_, row) in enumerate(detail_df.iterrows(), start=1):
        r = t2.add_row()
        _set_cell(r.cells[0], str(i), font_size=Pt(8))
        _set_cell(r.cells[1], str(row["title"])[:120], align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(8))
        decision = str(row["ai_tiab_decision"])
        _set_cell(r.cells[2], decision, font_size=Pt(8))
        # Highlight lost articles in red
        if row["ai_manteria"] == "Nao":
            _shade(r.cells[2], "FFD6D6")

    doc.add_paragraph()

    # --- Table 3: Lost articles (if any) ---
    if n_perdidos > 0:
        h3 = doc.add_heading("Table 3. Articles Missed by AI Screening", level=2)
        for r in h3.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(0, 0, 0)

        note = doc.add_paragraph(
            "These articles were included in the final systematic review but were "
            "excluded by the AI during TIAB screening. They would have been lost "
            "if AI screening was the sole method."
        )
        note.runs[0].font.size = Pt(9)
        note.runs[0].font.italic = True

        t3 = doc.add_table(rows=1, cols=3)
        t3.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_borders(t3)

        for j, txt in enumerate(["#", "Article Title", "AI Decision"]):
            _set_cell(t3.cell(0, j), txt, bold=True)
            _shade(t3.cell(0, j), "D9E2F3")

        for i, (_, row) in enumerate(lost_df.iterrows(), start=1):
            r = t3.add_row()
            _set_cell(r.cells[0], str(i), font_size=Pt(8))
            _set_cell(r.cells[1], str(row["title_display"])[:120],
                      align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(8))
            _set_cell(r.cells[2], str(row["ai_decision"]), font_size=Pt(8))

    # Save
    docx_path = output_dir / f"fulltext_check_{ts}.docx"
    doc.save(str(docx_path))

    # JSON summary
    summary = {
        "timestamp": ts,
        "total_fulltext": int(total_fulltext),
        "matched_in_ai": int(n_matched),
        "not_found": int(n_not_found),
        "ai_would_keep": int(n_passaria),
        "ai_would_miss": int(n_perdidos),
        "capture_rate": round(float(capture_rate), 4) if not np.isnan(capture_rate) else None,
        "miss_rate": round(float(miss_rate), 4) if not np.isnan(miss_rate) else None,
        "lost_titles": [str(row["title_display"]) for _, row in lost_df.iterrows()],
    }
    json_path = output_dir / f"fulltext_check_{ts}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)

    print(f"\n  Arquivos gerados em: {output_dir}")
    print(f"    - {xlsx_detail.name}")
    print(f"    - {docx_path.name}")
    print(f"    - {json_path.name}")
    print()

    return {
        "xlsx": xlsx_detail,
        "docx": docx_path,
        "json": json_path,
    }


# ----------------------------------------------- Word formatting helpers --

def _shade(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
              font_size=Pt(9), font_name="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)


def _add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


# ------------------------------------------------------------------ main --

def main():
    parser = argparse.ArgumentParser(
        description=(
            "Etapa 3 - Fulltext Check.\n\n"
            "Verifica se os artigos incluidos na revisao final (fulltext)\n"
            "teriam sido mantidos pela IA na triagem de TIAB.\n\n"
            "Modo automatico: coloque os arquivos em input/ e rode sem argumentos.\n"
            "Modo manual:     python diagnostic/03_fulltext_check.py --ai <arq> --fulltext <arq>"
        ),
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument("--ai", default=None,
                        help="Arquivo da IA (CSV/XLSX) com screening_decision.")
    parser.add_argument("--fulltext", default=None,
                        help="Arquivo com artigos incluidos na revisao final (CSV/XLSX).")
    args = parser.parse_args()

    ensure_folders()

    if args.ai and args.fulltext:
        ai_path, ft_path = args.ai, args.fulltext
    elif args.ai or args.fulltext:
        print("  ERRO: Informe ambos --ai e --fulltext, ou nenhum (auto-detectar).")
        sys.exit(1)
    else:
        print("\n  Modo automatico: buscando arquivos em input/ ...\n")
        ai_path, ft_path = auto_detect_files()

    for label, path in [("IA", ai_path), ("Fulltext", ft_path)]:
        if not os.path.isfile(path):
            print(f"  ERRO: arquivo {label} nao encontrado: {path}")
            sys.exit(1)

    print(f"  Arquivo IA:       {Path(ai_path).name}")
    print(f"  Arquivo Fulltext: {Path(ft_path).name}")

    run_fulltext_check(ai_path, ft_path, OUTPUT_DIR)


if __name__ == "__main__":
    main()

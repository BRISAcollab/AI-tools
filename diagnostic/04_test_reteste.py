"""
ETAPA 4 - Teste-Reteste (Test-Retest Reliability).

Compara duas execucoes da mesma IA (mesmo modelo) sobre o mesmo
banco de dados para avaliar a reprodutibilidade/consistencia
das decisoes de triagem.

Metricas: concordancia percentual, matriz de confusao IA1 vs IA2,
Cohen's Kappa, artigos discordantes detalhados.

Uso:
    1. Coloque os 2 arquivos da IA (1o e 2o teste) na pasta input/
    2. Rode:  python diagnostic/04_test_reteste.py

    Modo manual:
        python diagnostic/04_test_reteste.py --test1 input/t1.xlsx --test2 input/t2.xlsx
"""

import argparse
import sys
import os
import json
import datetime
from pathlib import Path

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ------------------------------------------------------------------ paths --
SCRIPT_DIR  = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
INPUT_DIR   = PROJECT_DIR / "input"
OUTPUT_DIR  = PROJECT_DIR / "output"


# ---------------------------------------------------------------- helpers --

def ensure_folders():
    INPUT_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)


def load_file(path: str) -> pd.DataFrame:
    ext = Path(path).suffix.lower()
    if ext == ".csv":
        try:
            return pd.read_csv(path, encoding="utf-8")
        except UnicodeDecodeError:
            return pd.read_csv(path, encoding="latin-1")
    elif ext in (".xlsx", ".xls"):
        return pd.read_excel(path)
    else:
        raise ValueError(f"Formato nao suportado: {ext}. Use .csv, .xlsx ou .xls")


def normalise_columns(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    return df


def normalise_title(s) -> str:
    if pd.isna(s):
        return ""
    return str(s).strip().lower()


def normalise_decision(s) -> str:
    """Padroniza variantes: included->include, excluded->exclude."""
    if pd.isna(s):
        return ""
    d = str(s).strip().lower()
    if d == "included":
        return "include"
    if d == "excluded":
        return "exclude"
    return d


def auto_detect_files():
    """
    Detecta automaticamente os dois arquivos da IA em input/.
    Ambos devem ter 'screening_decision' e quantidade similar de registros.
    """
    valid_ext = {".csv", ".xlsx", ".xls"}
    candidates = [f for f in INPUT_DIR.iterdir() if f.suffix.lower() in valid_ext]

    if len(candidates) < 2:
        print(f"\n  ERRO: Sao necessarios pelo menos 2 arquivos em '{INPUT_DIR}'.")
        sys.exit(1)

    # Collect metadata
    scored = []
    for f in candidates:
        try:
            df = normalise_columns(load_file(str(f)))
            cols = set(df.columns)
            n = len(df)
        except Exception:
            continue
        scored.append({
            "path": f,
            "n": n,
            "has_screening": "screening_decision" in cols,
            "has_id": "id" in cols,
            "name": f.name.lower(),
        })

    # Filter: only AI files (have screening_decision)
    ai_files = [s for s in scored if s["has_screening"]]
    ai_files.sort(key=lambda x: x["name"])

    if len(ai_files) < 2:
        print("  ERRO: Sao necessarios pelo menos 2 arquivos da IA (com screening_decision).")
        print("  Use: python diagnostic/04_test_reteste.py --test1 <arq1> --test2 <arq2>")
        sys.exit(1)

    # Try to identify by name pattern (1º teste, 2º teste)
    test1 = None
    test2 = None
    for s in ai_files:
        name = s["name"]
        if "1" in name and ("teste" in name or "test" in name):
            test1 = s
        elif "2" in name and ("teste" in name or "test" in name):
            test2 = s

    # Fallback: first two AI files sorted by name
    if test1 is None or test2 is None:
        test1 = ai_files[0]
        test2 = ai_files[1]

    return str(test1["path"]), str(test2["path"])


# ------------------------------------------------------------ statistics --

def binarise(decision: str) -> str:
    """include/included -> maybe, excluded -> exclude para binarizar."""
    d = normalise_decision(decision)
    if d == "include":
        return "maybe"
    return d


def confusion_matrix(df, col1, col2):
    """Matriz de confusao 2x2 entre duas colunas binarizadas."""
    a = df[col1]
    b = df[col2]
    tp = int(((a == "maybe")  & (b == "maybe")).sum())
    fp = int(((a == "maybe")  & (b == "exclude")).sum())
    fn = int(((a == "exclude") & (b == "maybe")).sum())
    tn = int(((a == "exclude") & (b == "exclude")).sum())
    return tp, fp, fn, tn


def calc_kappa(tp, fp, fn, tn):
    n = tp + fp + fn + tn
    if n == 0:
        return float("nan"), float("nan"), float("nan"), float("nan"), ""
    po = (tp + tn) / n
    pe = ((tp + fp) * (tp + fn) + (fn + tn) * (fp + tn)) / (n * n)
    if pe == 1:
        k = 1.0
    else:
        k = (po - pe) / (1 - pe)
    if (1 - pe) == 0:
        se = float("nan")
    else:
        se = np.sqrt(pe * (1 - pe) / (n * (1 - pe) ** 2))
    ci_lo = k - 1.96 * se
    ci_hi = k + 1.96 * se
    if k < 0:       interp = "Poor (< 0)"
    elif k < 0.20:  interp = "Slight (0.00-0.20)"
    elif k < 0.40:  interp = "Fair (0.21-0.40)"
    elif k < 0.60:  interp = "Moderate (0.41-0.60)"
    elif k < 0.80:  interp = "Substantial (0.61-0.80)"
    else:            interp = "Almost Perfect (0.81-1.00)"
    return k, se, ci_lo, ci_hi, interp


def fmt(v, d=4):
    if isinstance(v, int):
        return str(v)
    if isinstance(v, float):
        if np.isnan(v):
            return "N/A"
        if np.isinf(v):
            return "-"
        return f"{v:.{d}f}"
    return str(v)


def fmt_pct(v):
    if isinstance(v, float) and not np.isnan(v) and not np.isinf(v):
        return f"{v * 100:.1f}%"
    return "-"


# --------------------------------------------------------------- analysis --

def run_test_retest(test1_path: str, test2_path: str, output_dir: Path):
    """Compara duas execucoes da IA (test-retest)."""

    t1_df = normalise_columns(load_file(test1_path))
    t2_df = normalise_columns(load_file(test2_path))

    # Detectar coluna de decisao
    decision_col = "screening_decision"
    for df, label in [(t1_df, "Teste 1"), (t2_df, "Teste 2")]:
        if decision_col not in df.columns:
            raise KeyError(f"Arquivo {label} sem coluna '{decision_col}'. Colunas: {list(df.columns)}")
        if "title" not in df.columns:
            raise KeyError(f"Arquivo {label} sem coluna 'title'. Colunas: {list(df.columns)}")

    # Normalizar titulos e preparar merge key com cumcount
    for df in (t1_df, t2_df):
        df["_title_key"] = df["title"].apply(normalise_title)
        df["_occ"] = df.groupby("_title_key").cumcount().astype(str)
        df["_merge_key"] = df["_title_key"] + "__" + df["_occ"]

    # Merge
    merged = pd.merge(
        t1_df[["title", "_merge_key", decision_col]],
        t2_df[["_merge_key", decision_col]],
        on="_merge_key",
        how="outer",
        suffixes=("_t1", "_t2"),
    )

    col_t1 = f"{decision_col}_t1"
    col_t2 = f"{decision_col}_t2"

    n_total = len(merged)
    n_matched = merged[[col_t1, col_t2]].notna().all(axis=1).sum()

    # Binarizar
    merged["t1_bin"] = merged[col_t1].apply(binarise)
    merged["t2_bin"] = merged[col_t2].apply(binarise)

    # Concordancia por classe original (antes de binarizar)
    merged["t1_orig"] = merged[col_t1].astype(str).str.strip().str.lower()
    merged["t2_orig"] = merged[col_t2].astype(str).str.strip().str.lower()

    # ================================================================
    #  CONCORDANCIA GRANULAR (3 categorias: include, maybe, exclude)
    # ================================================================
    exact_match = (merged["t1_orig"] == merged["t2_orig"]).sum()
    exact_pct = exact_match / n_total if n_total > 0 else float("nan")

    # Crosstab 3x3
    categories = ["include", "maybe", "exclude"]
    cross_3x3 = pd.crosstab(
        merged["t1_orig"].apply(lambda x: x if x in categories else "other"),
        merged["t2_orig"].apply(lambda x: x if x in categories else "other"),
        margins=True,
    )

    # ================================================================
    #  CONCORDANCIA BINARIZADA (maybe vs exclude)
    # ================================================================
    binary_match = (merged["t1_bin"] == merged["t2_bin"]).sum()
    binary_pct = binary_match / n_total if n_total > 0 else float("nan")

    tp, fp, fn, tn = confusion_matrix(merged, "t1_bin", "t2_bin")
    kappa, se, ci_lo, ci_hi, interp = calc_kappa(tp, fp, fn, tn)

    # Discordancias
    disc_df = merged[merged["t1_bin"] != merged["t2_bin"]].copy()
    disc_df["title_display"] = disc_df["title"].fillna("")

    # Classificar tipo de discordancia
    disc_df["tipo"] = disc_df.apply(
        lambda r: f"{r['t1_orig']} -> {r['t2_orig']}", axis=1
    )

    # ================================================================
    #  CONSOLE
    # ================================================================
    t1_name = Path(test1_path).stem
    t2_name = Path(test2_path).stem

    print()
    print("=" * 64)
    print("  ETAPA 4 - TESTE-RETESTE (Test-Retest Reliability)")
    print("=" * 64)

    print(f"\n  Teste 1: {t1_name}")
    print(f"  Teste 2: {t2_name}")
    print(f"  Total pareados: {n_total}")

    print(f"\n  CONCORDANCIA EXATA (3 categorias)")
    print("-" * 64)
    print(f"  Concordantes:  {exact_match}/{n_total}  ({exact_pct*100:.1f}%)")
    print(f"  Discordantes:  {n_total - exact_match}/{n_total}  ({(1 - exact_pct)*100:.1f}%)")

    print(f"\n  TABELA CRUZADA (Teste 1 x Teste 2)")
    print("-" * 64)
    # Print cross table nicely
    print(f"{'':>16}", end="")
    for col in cross_3x3.columns:
        print(f"  {str(col):>10}", end="")
    print()
    for idx in cross_3x3.index:
        print(f"  {str(idx):>14}", end="")
        for col in cross_3x3.columns:
            val = cross_3x3.loc[idx, col]
            print(f"  {val:>10}", end="")
        print()

    print(f"\n  CONCORDANCIA BINARIZADA (maybe/include vs exclude)")
    print("-" * 64)
    print(f"  Concordantes:  {binary_match}/{n_total}  ({binary_pct*100:.1f}%)")
    print(f"  Discordantes:  {n_total - binary_match}/{n_total}  ({(1 - binary_pct)*100:.1f}%)")

    print(f"\n  MATRIZ DE CONFUSAO (binarizada)")
    print("-" * 64)
    print(f"{'':>20}{'Teste 2':^28}")
    print(f"{'':>20}{'maybe':>12}{'exclude':>12}")
    print(f"  {'T1  maybe':>18}{tp:>12}{fp:>12}   | {tp + fp}")
    print(f"  {'    exclude':>18}{fn:>12}{tn:>12}   | {fn + tn}")
    print(f"{'':>20}{'-'*12:>12}{'-'*12:>12}")
    print(f"{'':>20}{tp + fn:>12}{fp + tn:>12}{'':>5}{n_total}")

    print(f"\n  KAPPA (Test-Retest)")
    print("-" * 64)
    print(f"  Kappa:         {fmt(kappa)}")
    print(f"  SE:            {fmt(se)}")
    print(f"  95% CI:        [{fmt(ci_lo)}, {fmt(ci_hi)}]")
    print(f"  Interpretacao: {interp}")

    if len(disc_df) > 0:
        print(f"\n  ARTIGOS DISCORDANTES ({len(disc_df)})")
        print("-" * 64)
        tipo_counts = disc_df["tipo"].value_counts()
        for tipo, count in tipo_counts.items():
            print(f"    {tipo}: {count} artigo(s)")
        if len(disc_df) <= 50:
            print()
            for _, row in disc_df.iterrows():
                t = str(row["title_display"])[:80]
                print(f"    [{row['t1_orig']} -> {row['t2_orig']}] {t}")
    else:
        print(f"\n  PERFEITO: As duas execucoes da IA concordaram em TODOS os {n_total} artigos!")

    print("=" * 64)

    # ================================================================
    #  GERAR XLSX
    # ================================================================
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    detail_df = merged[["title", "t1_orig", "t2_orig", "t1_bin", "t2_bin"]].copy()
    detail_df.columns = ["title", "test1_original", "test2_original",
                         "test1_binary", "test2_binary"]
    detail_df["concordante_exato"] = (detail_df["test1_original"] == detail_df["test2_original"]).map(
        {True: "Sim", False: "Nao"}
    )
    detail_df["concordante_binario"] = (detail_df["test1_binary"] == detail_df["test2_binary"]).map(
        {True: "Sim", False: "Nao"}
    )
    detail_df.insert(0, "n", range(1, len(detail_df) + 1))

    xlsx_path = output_dir / f"test_reteste_{ts}.xlsx"
    with pd.ExcelWriter(xlsx_path, engine="openpyxl") as writer:
        detail_df.to_excel(writer, sheet_name="Detalhamento", index=False)
        cross_3x3.to_excel(writer, sheet_name="Crosstab_3x3")
        if len(disc_df) > 0:
            disc_out = disc_df[["title_display", "t1_orig", "t2_orig", "tipo"]].copy()
            disc_out.columns = ["title", "test1", "test2", "tipo_discordancia"]
            disc_out.to_excel(writer, sheet_name="Discordantes", index=False)

    # ================================================================
    #  GERAR WORD
    # ================================================================
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Test-Retest Reliability Analysis - AI TIAB Screening")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run(f"Test 1: {t1_name}\nTest 2: {t2_name}")
    run.font.size = Pt(9)
    run.font.name = "Times New Roman"

    # --- Table 1: Summary ---
    _add_heading(doc, "Table 1. Test-Retest Agreement Summary")

    t1_tbl = doc.add_table(rows=1, cols=2)
    t1_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(t1_tbl)
    for j, txt in enumerate(["Parameter", "Value"]):
        _set_cell(t1_tbl.cell(0, j), txt, bold=True)
        _shade(t1_tbl.cell(0, j), "D9E2F3")

    summary_rows = [
        ("Total articles compared",                 str(n_total)),
        ("Exact agreement (3 categories)",           f"{exact_match} ({exact_pct*100:.1f}%)"),
        ("Binary agreement (maybe/include vs exc.)", f"{binary_match} ({binary_pct*100:.1f}%)"),
        ("Discordant articles (exact)",              str(n_total - exact_match)),
        ("Discordant articles (binary)",             str(n_total - binary_match)),
        ("Cohen's Kappa",                            f"{fmt(kappa)} [{fmt(ci_lo)} - {fmt(ci_hi)}]"),
        ("Kappa Interpretation",                     interp),
    ]
    for label, val in summary_rows:
        row = t1_tbl.add_row()
        _set_cell(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row.cells[1], val)

    doc.add_paragraph()

    # --- Table 2: Cross-tabulation 3x3 ---
    _add_heading(doc, "Table 2. Cross-Tabulation of Decisions (3 Categories)")

    note = doc.add_paragraph(
        "Rows represent Test 1 decisions, columns represent Test 2 decisions. "
        "Values on the diagonal indicate agreement."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.italic = True

    cat_labels = [c for c in cross_3x3.columns if c != "All"]
    cat_labels_with_total = list(cross_3x3.columns)
    idx_labels_with_total = list(cross_3x3.index)

    n_cols = len(cat_labels_with_total) + 1  # +1 for row label
    n_rows = len(idx_labels_with_total) + 1  # +1 for header

    t2_tbl = doc.add_table(rows=n_rows, cols=n_cols)
    t2_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(t2_tbl)

    _set_cell(t2_tbl.cell(0, 0), "Test 1 \\ Test 2", bold=True, font_size=Pt(8))
    _shade(t2_tbl.cell(0, 0), "D9E2F3")
    for j, col in enumerate(cat_labels_with_total):
        _set_cell(t2_tbl.cell(0, j + 1), str(col), bold=True, font_size=Pt(8))
        _shade(t2_tbl.cell(0, j + 1), "D9E2F3")

    for i, idx in enumerate(idx_labels_with_total):
        _set_cell(t2_tbl.cell(i + 1, 0), str(idx), bold=(str(idx) == "All"),
                  align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(8))
        if str(idx) == "All":
            _shade(t2_tbl.cell(i + 1, 0), "E8E8E8")
        for j, col in enumerate(cat_labels_with_total):
            val = cross_3x3.loc[idx, col]
            cell = t2_tbl.cell(i + 1, j + 1)
            _set_cell(cell, str(val), font_size=Pt(8),
                      bold=(str(idx) == "All" or str(col) == "All"))
            # Highlight diagonal (agreement)
            if str(idx) == str(col) and str(idx) != "All":
                _shade(cell, "D5F5E3")
            elif str(idx) == "All" or str(col) == "All":
                _shade(cell, "E8E8E8")

    doc.add_paragraph()

    # --- Table 3: Confusion Matrix (binary) ---
    _add_heading(doc, "Table 3. Binary Confusion Matrix (Test 1 vs Test 2)")

    note2 = doc.add_paragraph(
        "Decisions binarised: include and maybe grouped as positive (pass screening), "
        "exclude as negative."
    )
    note2.runs[0].font.size = Pt(9)
    note2.runs[0].font.italic = True

    t3_tbl = doc.add_table(rows=4, cols=4)
    t3_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(t3_tbl)

    # Header row
    _set_cell(t3_tbl.cell(0, 0), "", bold=True)
    _set_cell(t3_tbl.cell(0, 1), "", bold=True)
    _set_cell(t3_tbl.cell(0, 2), "Test 2", bold=True)
    _shade(t3_tbl.cell(0, 2), "D9E2F3")
    _set_cell(t3_tbl.cell(0, 3), "", bold=True)
    t3_tbl.cell(0, 2).merge(t3_tbl.cell(0, 3))

    _set_cell(t3_tbl.cell(1, 0), "", bold=True)
    _set_cell(t3_tbl.cell(1, 1), "", bold=True)
    _set_cell(t3_tbl.cell(1, 2), "maybe", bold=True)
    _shade(t3_tbl.cell(1, 2), "D9E2F3")
    _set_cell(t3_tbl.cell(1, 3), "exclude", bold=True)
    _shade(t3_tbl.cell(1, 3), "D9E2F3")

    _set_cell(t3_tbl.cell(2, 0), "Test 1", bold=True)
    _shade(t3_tbl.cell(2, 0), "D9E2F3")
    _set_cell(t3_tbl.cell(2, 1), "maybe", bold=True)
    _shade(t3_tbl.cell(2, 1), "D9E2F3")
    _set_cell(t3_tbl.cell(2, 2), str(tp))
    _shade(t3_tbl.cell(2, 2), "D5F5E3")
    _set_cell(t3_tbl.cell(2, 3), str(fp))

    _set_cell(t3_tbl.cell(3, 0), "")
    _set_cell(t3_tbl.cell(3, 1), "exclude", bold=True)
    _shade(t3_tbl.cell(3, 1), "D9E2F3")
    _set_cell(t3_tbl.cell(3, 2), str(fn))
    _set_cell(t3_tbl.cell(3, 3), str(tn))
    _shade(t3_tbl.cell(3, 3), "D5F5E3")

    doc.add_paragraph()

    # --- Table 4: Kappa ---
    _add_heading(doc, "Table 4. Cohen's Kappa - Test-Retest Reliability")

    t4_tbl = doc.add_table(rows=1, cols=2)
    t4_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _add_borders(t4_tbl)
    for j, txt in enumerate(["Statistic", "Value"]):
        _set_cell(t4_tbl.cell(0, j), txt, bold=True)
        _shade(t4_tbl.cell(0, j), "D9E2F3")

    kappa_rows = [
        ("Observed Agreement (Po)",      fmt_pct(binary_match / n_total if n_total else float("nan"))),
        ("Cohen's Kappa",                fmt(kappa)),
        ("Standard Error",               fmt(se)),
        ("95% CI",                        f"[{fmt(ci_lo)}, {fmt(ci_hi)}]"),
        ("Interpretation (Landis & Koch)", interp),
    ]
    for label, val in kappa_rows:
        row = t4_tbl.add_row()
        _set_cell(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell(row.cells[1], val)

    doc.add_paragraph()

    # --- Table 5: Discordant articles (if not too many) ---
    if 0 < len(disc_df) <= 100:
        _add_heading(doc, "Table 5. Discordant Articles")

        note3 = doc.add_paragraph(
            "Articles where the two AI runs produced different binary decisions."
        )
        note3.runs[0].font.size = Pt(9)
        note3.runs[0].font.italic = True

        t5_tbl = doc.add_table(rows=1, cols=4)
        t5_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _add_borders(t5_tbl)
        for j, txt in enumerate(["#", "Article Title", "Test 1", "Test 2"]):
            _set_cell(t5_tbl.cell(0, j), txt, bold=True, font_size=Pt(8))
            _shade(t5_tbl.cell(0, j), "D9E2F3")

        for i, (_, row) in enumerate(disc_df.iterrows(), start=1):
            r = t5_tbl.add_row()
            _set_cell(r.cells[0], str(i), font_size=Pt(8))
            _set_cell(r.cells[1], str(row["title_display"])[:120],
                      align=WD_ALIGN_PARAGRAPH.LEFT, font_size=Pt(8))
            _set_cell(r.cells[2], str(row["t1_orig"]), font_size=Pt(8))
            _set_cell(r.cells[3], str(row["t2_orig"]), font_size=Pt(8))
            _shade(r.cells[2], "FFD6D6")
            _shade(r.cells[3], "FFD6D6")

    # Save Word
    docx_path = output_dir / f"test_reteste_{ts}.docx"
    doc.save(str(docx_path))

    # ================================================================
    #  JSON
    # ================================================================
    summary = {
        "timestamp": ts,
        "test1": Path(test1_path).name,
        "test2": Path(test2_path).name,
        "total": int(n_total),
        "exact_agreement": int(exact_match),
        "exact_agreement_pct": round(float(exact_pct), 4) if not np.isnan(exact_pct) else None,
        "binary_agreement": int(binary_match),
        "binary_agreement_pct": round(float(binary_pct), 4) if not np.isnan(binary_pct) else None,
        "confusion_matrix": {"tp": tp, "fp": fp, "fn": fn, "tn": tn},
        "kappa": round(float(kappa), 4) if not np.isnan(kappa) else None,
        "kappa_se": round(float(se), 4) if not np.isnan(se) else None,
        "kappa_ci": [round(float(ci_lo), 4), round(float(ci_hi), 4)],
        "kappa_interpretation": interp,
        "n_discordant_exact": int(n_total - exact_match),
        "n_discordant_binary": int(len(disc_df)),
    }
    json_path = output_dir / f"test_reteste_{ts}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)

    print(f"\n  Arquivos gerados em: {output_dir}")
    print(f"    - {xlsx_path.name}")
    print(f"    - {docx_path.name}")
    print(f"    - {json_path.name}")
    print()


# ----------------------------------------------- Word formatting helpers --

def _shade(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
              font_size=Pt(9), font_name="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)


def _add_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def _add_heading(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)


# ------------------------------------------------------------------ main --

def main():
    parser = argparse.ArgumentParser(
        description=(
            "Etapa 4 - Teste-Reteste (Test-Retest Reliability).\n\n"
            "Compara duas execucoes da IA no mesmo banco de dados\n"
            "para avaliar a reprodutibilidade das decisoes.\n\n"
            "Modo automatico: coloque os 2 arquivos da IA em input/.\n"
            "Modo manual:     python diagnostic/04_test_reteste.py --test1 <arq1> --test2 <arq2>"
        ),
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument("--test1", default=None,
                        help="Arquivo da IA - 1o teste (CSV/XLSX).")
    parser.add_argument("--test2", default=None,
                        help="Arquivo da IA - 2o teste (CSV/XLSX).")
    args = parser.parse_args()

    ensure_folders()

    if args.test1 and args.test2:
        t1_path, t2_path = args.test1, args.test2
    elif args.test1 or args.test2:
        print("  ERRO: Informe ambos --test1 e --test2, ou nenhum (auto-detectar).")
        sys.exit(1)
    else:
        print("\n  Modo automatico: buscando arquivos da IA em input/ ...\n")
        t1_path, t2_path = auto_detect_files()

    for label, path in [("Teste 1", t1_path), ("Teste 2", t2_path)]:
        if not os.path.isfile(path):
            print(f"  ERRO: arquivo {label} nao encontrado: {path}")
            sys.exit(1)

    print(f"  Teste 1: {Path(t1_path).name}")
    print(f"  Teste 2: {Path(t2_path).name}")

    run_test_retest(t1_path, t2_path, OUTPUT_DIR)


if __name__ == "__main__":
    main()

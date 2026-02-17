"""
ETAPA 2 - Analise Diagnostica e Cohen's Kappa.

Le o arquivo pareado gerado pela Etapa 1 (output/pareamento.xlsx),
calcula metricas diagnosticas e gera tabelas em Word (.docx)
prontas para publicacao.

Uso:
    python 02_analise_diagnostica.py
    python 02_analise_diagnostica.py --input output/pareamento.xlsx
"""

import argparse
import sys
import datetime
import json
from pathlib import Path

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ------------------------------------------------------------------ paths --
SCRIPT_DIR  = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent
OUTPUT_DIR  = PROJECT_DIR / "output"
DEFAULT_INPUT = OUTPUT_DIR / "pareamento.xlsx"


# ============================================================ statistics ==

def normalise_decision(s) -> str:
    """Padroniza variantes: included->include, excluded->exclude."""
    if pd.isna(s):
        return ""
    d = str(s).strip().lower()
    if d == "included":
        return "include"
    if d == "excluded":
        return "exclude"
    return d


def binarise(df: pd.DataFrame) -> pd.DataFrame:
    """include/included -> maybe, excluded -> exclude para binarizar."""
    df = df.copy()
    df["screening_decision"] = (
        df["screening_decision"]
        .apply(normalise_decision)
        .replace("include", "maybe")
    )
    df["decision_human"] = (
        df["decision_human"]
        .apply(normalise_decision)
        .replace("include", "maybe")
    )
    return df


def confusion(df):
    tp = int(((df["screening_decision"] == "maybe")  & (df["decision_human"] == "maybe")).sum())
    fp = int(((df["screening_decision"] == "maybe")  & (df["decision_human"] == "exclude")).sum())
    fn = int(((df["screening_decision"] == "exclude") & (df["decision_human"] == "maybe")).sum())
    tn = int(((df["screening_decision"] == "exclude") & (df["decision_human"] == "exclude")).sum())
    return tp, fp, fn, tn


def calc_metrics(tp, fp, fn, tn):
    n    = tp + fp + fn + tn
    sens = tp / (tp + fn) if (tp + fn) else float("nan")
    spec = tn / (tn + fp) if (tn + fp) else float("nan")
    ppv  = tp / (tp + fp) if (tp + fp) else float("nan")
    npv  = tn / (tn + fn) if (tn + fn) else float("nan")
    acc  = (tp + tn) / n  if n else float("nan")
    prev = (tp + fn) / n  if n else float("nan")
    lr_p = sens / (1 - spec) if (1 - spec) > 0 else float("inf")
    lr_n = (1 - sens) / spec if spec > 0 else float("inf")
    f1   = 2*tp / (2*tp + fp + fn) if (2*tp + fp + fn) else float("nan")
    youd = sens + spec - 1
    return {
        "N":                          n,
        "Prevalence":                 prev,
        "Sensitivity (Recall)":       sens,
        "Specificity":                spec,
        "PPV (Precision)":            ppv,
        "NPV":                        npv,
        "Accuracy":                   acc,
        "F1 Score":                   f1,
        "LR+":                        lr_p,
        "LR-":                        lr_n,
        "Youden Index":               youd,
    }


def calc_kappa(tp, fp, fn, tn):
    n = tp + fp + fn + tn
    if n == 0:
        return float("nan"), float("nan"), float("nan"), float("nan"), ""
    po = (tp + tn) / n
    pe = ((tp+fp)*(tp+fn) + (fn+tn)*(fp+tn)) / (n*n)
    if pe == 1:
        k = 1.0
    else:
        k = (po - pe) / (1 - pe)
    # SE
    if (1 - pe) == 0:
        se = float("nan")
    else:
        se = np.sqrt(pe * (1 - pe) / (n * (1 - pe)**2))
    ci_lo = k - 1.96 * se
    ci_hi = k + 1.96 * se
    # interpretation
    if k < 0:       interp = "Poor (< 0)"
    elif k < 0.20:  interp = "Slight (0.00-0.20)"
    elif k < 0.40:  interp = "Fair (0.21-0.40)"
    elif k < 0.60:  interp = "Moderate (0.41-0.60)"
    elif k < 0.80:  interp = "Substantial (0.61-0.80)"
    else:            interp = "Almost Perfect (0.81-1.00)"
    return k, se, ci_lo, ci_hi, interp


# ======================================================== formatting ==

def fmt(v, d=4):
    if isinstance(v, int):   return str(v)
    if isinstance(v, float):
        if np.isnan(v): return "N/A"
        if np.isinf(v): return "-"
        return f"{v:.{d}f}"
    return str(v)


def fmt_pct(v):
    if isinstance(v, float) and not np.isnan(v) and not np.isinf(v):
        return f"{v*100:.1f}"
    return "-"


def fmt_ci(lo, hi):
    return f"{fmt(lo)} - {fmt(hi)}"


# ======================================================= Word helpers ==

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=Pt(9), font_name="Times New Roman"):
    """Set formatted text inside a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.size = font_size
    run.font.name = font_name
    run.bold = bold
    # Spacing
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after  = Pt(1)


def add_table_borders(table):
    """Add thin borders to all cells of a Word table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def add_heading(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


# ========================================== Word document generation ==

def generate_word(df, tp, fp, fn, tn, metrics, kappa_val, se_val,
                  ci_lo, ci_hi, interp, output_dir: Path):
    """Generate a publication-ready Word document with result tables."""

    doc = Document()

    # --- Document style defaults ---
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ================================================================
    #  TITLE
    # ================================================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Diagnostic Test Results - AI vs Human Screening")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"

    ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(f"Generated: {ts}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    n = tp + fp + fn + tn

    # ================================================================
    #  TABLE 1 - Sample Characteristics
    # ================================================================
    add_heading(doc, "Table 1. Sample Characteristics")

    t1 = doc.add_table(rows=5, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t1)

    t1_data = [
        ("Characteristic",                  "Value"),
        ("Total articles analyzed, n",      str(n)),
        ("AI positive (maybe/include), n (%)",
         f"{tp+fp} ({(tp+fp)/n*100:.1f}%)" if n else "-"),
        ("Human positive (maybe), n (%)",
         f"{tp+fn} ({(tp+fn)/n*100:.1f}%)" if n else "-"),
        ("Prevalence of human positive",
         f"{fmt(metrics['Prevalence'])} ({fmt_pct(metrics['Prevalence'])}%)"),
    ]
    for i, (left, right) in enumerate(t1_data):
        is_header = (i == 0)
        set_cell_text(t1.cell(i, 0), left, bold=is_header,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(t1.cell(i, 1), right, bold=is_header)
        if is_header:
            set_cell_shading(t1.cell(i, 0), "D9E2F3")
            set_cell_shading(t1.cell(i, 1), "D9E2F3")

    doc.add_paragraph()

    # ================================================================
    #  TABLE 2 - Confusion Matrix (2x2)
    # ================================================================
    add_heading(doc, "Table 2. Confusion Matrix")

    note = doc.add_paragraph(
        "Gold standard: Human decision. "
        "Positive: maybe (article passes screening). "
        "Negative: exclude."
    )
    note.runs[0].font.size = Pt(8)
    note.runs[0].font.italic = True

    t2 = doc.add_table(rows=4, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t2)

    # Header row
    headers2 = ["", "Human: Maybe", "Human: Exclude", "Total"]
    for j, h in enumerate(headers2):
        set_cell_text(t2.cell(0, j), h, bold=True)
        set_cell_shading(t2.cell(0, j), "D9E2F3")

    # Data rows
    cm_rows = [
        ("AI: Maybe",   str(tp),    str(fp),    str(tp + fp)),
        ("AI: Exclude", str(fn),    str(tn),    str(fn + tn)),
        ("Total",       str(tp+fn), str(fp+tn), str(n)),
    ]
    for i, (label, *vals) in enumerate(cm_rows, start=1):
        set_cell_text(t2.cell(i, 0), label, bold=True,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
        for j, v in enumerate(vals, start=1):
            set_cell_text(t2.cell(i, j), v)
        if i == 3:  # total row
            for j in range(4):
                set_cell_text(t2.cell(i, j),
                              t2.cell(i, j).text, bold=True)

    doc.add_paragraph()

    # ================================================================
    #  TABLE 3 - Diagnostic Accuracy
    # ================================================================
    add_heading(doc, "Table 3. Diagnostic Accuracy of AI Screening")

    t3 = doc.add_table(rows=1, cols=3)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t3)

    # Header
    for j, h in enumerate(["Metric", "Value", "% Value"]):
        set_cell_text(t3.cell(0, j), h, bold=True)
        set_cell_shading(t3.cell(0, j), "D9E2F3")

    metrics_display = [
        ("Sensitivity (Recall)", metrics["Sensitivity (Recall)"]),
        ("Specificity",          metrics["Specificity"]),
        ("Positive Predictive Value (Precision)", metrics["PPV (Precision)"]),
        ("Negative Predictive Value",             metrics["NPV"]),
        ("Accuracy",             metrics["Accuracy"]),
        ("F1 Score",             metrics["F1 Score"]),
        ("Positive Likelihood Ratio (LR+)", metrics["LR+"]),
        ("Negative Likelihood Ratio (LR-)", metrics["LR-"]),
        ("Youden Index",         metrics["Youden Index"]),
    ]
    for name, val in metrics_display:
        row = t3.add_row()
        set_cell_text(row.cells[0], name, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[1], fmt(val))
        pct = fmt_pct(val) + "%" if fmt_pct(val) != "-" else "-"
        set_cell_text(row.cells[2], pct)

    doc.add_paragraph()

    # ================================================================
    #  TABLE 4 - Inter-rater Agreement (Cohen's Kappa)
    # ================================================================
    add_heading(doc, "Table 4. Inter-rater Agreement (Cohen's Kappa)")

    t4 = doc.add_table(rows=1, cols=2)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t4)

    for j, h in enumerate(["Parameter", "Value"]):
        set_cell_text(t4.cell(0, j), h, bold=True)
        set_cell_shading(t4.cell(0, j), "D9E2F3")

    po = (tp + tn) / n if n else float("nan")
    pe = ((tp+fp)*(tp+fn) + (fn+tn)*(fp+tn)) / (n*n) if n else float("nan")

    kappa_rows = [
        ("Observed agreement (Po)",  fmt(po)),
        ("Expected agreement (Pe)",  fmt(pe)),
        ("Cohen's Kappa",            fmt(kappa_val)),
        ("Standard Error (SE)",      fmt(se_val)),
        ("95% CI",                   f"[{fmt(ci_lo)}, {fmt(ci_hi)}]"),
        ("Interpretation (Landis & Koch)", interp),
    ]
    for label, val in kappa_rows:
        row = t4.add_row()
        set_cell_text(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[1], val)

    # Footnote
    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Interpretation: < 0 Poor; 0.00-0.20 Slight; "
        "0.21-0.40 Fair; 0.41-0.60 Moderate; "
        "0.61-0.80 Substantial; 0.81-1.00 Almost Perfect "
        "(Landis & Koch, 1977)."
    )
    foot.runs[0].font.size = Pt(8)
    foot.runs[0].font.italic = True

    # ================================================================
    #  TABLE 5 - Summary for abstract / results section
    # ================================================================
    add_heading(doc, "Table 5. Summary of Results")

    t5 = doc.add_table(rows=1, cols=2)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(t5)

    for j, h in enumerate(["", ""]):
        set_cell_shading(t5.cell(0, j), "D9E2F3")
    set_cell_text(t5.cell(0, 0), "Outcome", bold=True,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(t5.cell(0, 1), "Result", bold=True)

    summary_rows = [
        ("Articles screened",                              str(n)),
        ("True Positives",                                 str(tp)),
        ("False Positives",                                str(fp)),
        ("False Negatives",                                str(fn)),
        ("True Negatives",                                 str(tn)),
        ("Sensitivity",                f"{fmt_pct(metrics['Sensitivity (Recall)'])}%"),
        ("Specificity",                f"{fmt_pct(metrics['Specificity'])}%"),
        ("PPV",                        f"{fmt_pct(metrics['PPV (Precision)'])}%"),
        ("NPV",                        f"{fmt_pct(metrics['NPV'])}%"),
        ("Accuracy",                   f"{fmt_pct(metrics['Accuracy'])}%"),
        ("Cohen's Kappa (95% CI)",     f"{fmt(kappa_val)} [{fmt(ci_lo)}, {fmt(ci_hi)}]"),
        ("Agreement",                  interp),
    ]
    for label, val in summary_rows:
        row = t5.add_row()
        set_cell_text(row.cells[0], label, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[1], val)

    # ================================================================
    #  SAVE
    # ================================================================
    ts_file = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = output_dir / f"diagnostic_results_{ts_file}.docx"
    doc.save(str(docx_path))
    return docx_path


# ========================================== console report (brief) ==

def print_console(tp, fp, fn, tn, metrics, k, se, ci_lo, ci_hi, interp):
    n = tp + fp + fn + tn
    print()
    print("=" * 62)
    print("  ETAPA 2 - ANALISE DIAGNOSTICA")
    print("=" * 62)

    print(f"\n  MATRIZ DE CONFUSAO  (gold standard = humano)")
    print("-" * 62)
    print(f"                        Humano")
    print(f"                    maybe    exclude")
    print(f"  IA  maybe      {tp:>6}    {fp:>6}     | {tp+fp}")
    print(f"      exclude    {fn:>6}    {tn:>6}     | {fn+tn}")
    print(f"                 ------    ------")
    print(f"                 {tp+fn:>6}    {fp+tn:>6}       {n}")

    print(f"\n  METRICAS")
    print("-" * 62)
    for name, val in metrics.items():
        pct = fmt_pct(val)
        pct_s = f"  ({pct}%)" if pct != "-" else ""
        print(f"  {name:<35s}  {fmt(val)}{pct_s}")

    print(f"\n  KAPPA")
    print("-" * 62)
    print(f"  Kappa:         {fmt(k)}")
    print(f"  SE:            {fmt(se)}")
    print(f"  95% CI:        [{fmt(ci_lo)}, {fmt(ci_hi)}]")
    print(f"  Interpretacao: {interp}")
    print("=" * 62)


# ------------------------------------------------------------------ main --

def main():
    parser = argparse.ArgumentParser(
        description=(
            "Etapa 2 - Analise diagnostica.\n\n"
            "Le o arquivo pareado (output/pareamento.xlsx) gerado pela Etapa 1\n"
            "e gera tabelas em Word prontas para publicacao."
        ),
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument("--input", "-i", default=None,
                        help="Arquivo pareado XLSX/CSV (padrao: output/pareamento.xlsx).")
    args = parser.parse_args()

    input_path = Path(args.input) if args.input else DEFAULT_INPUT

    if not input_path.is_file():
        print(f"\n  ERRO: Arquivo pareado nao encontrado: {input_path}")
        print("  Rode primeiro: python 01_pareamento.py\n")
        sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Load paired data
    ext = input_path.suffix.lower()
    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(str(input_path))
    else:
        df = pd.read_csv(str(input_path), encoding="utf-8-sig")

    # Normalise columns
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    if "screening_decision" not in df.columns or "decision_human" not in df.columns:
        print("  ERRO: Colunas 'screening_decision' e 'decision_human' nao encontradas.")
        print(f"  Colunas do arquivo: {list(df.columns)}")
        sys.exit(1)

    print(f"\n  Arquivo lido: {input_path.name}  ({len(df)} registros)")

    # Binarise
    df = binarise(df)

    # Filter valid
    valid = df["screening_decision"].isin(["maybe", "exclude"]) & \
            df["decision_human"].isin(["maybe", "exclude"])
    inv = (~valid).sum()
    if inv:
        print(f"  {inv} registro(s) com decisao invalida removido(s).")
    df = df[valid].copy()

    if df.empty:
        print("  ERRO: Nenhum registro valido.")
        sys.exit(1)

    # Calculate
    tp, fp, fn, tn = confusion(df)
    metrics = calc_metrics(tp, fp, fn, tn)
    k, se, ci_lo, ci_hi, interp = calc_kappa(tp, fp, fn, tn)

    # Console
    print_console(tp, fp, fn, tn, metrics, k, se, ci_lo, ci_hi, interp)

    # Word document
    docx_path = generate_word(df, tp, fp, fn, tn, metrics, k, se,
                              ci_lo, ci_hi, interp, OUTPUT_DIR)

    print(f"\n  Documento Word gerado: {docx_path.name}")
    print(f"  Pasta: {OUTPUT_DIR}\n")

    # Also save JSON
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    summary = {
        "n": tp + fp + fn + tn,
        "confusion_matrix": {"TP": tp, "FP": fp, "FN": fn, "TN": tn},
        "metrics": {name: (round(v, 4) if isinstance(v, float) and not np.isnan(v) else v)
                    for name, v in metrics.items()},
        "kappa": round(k, 4),
        "kappa_se": round(se, 4),
        "kappa_ci_95": [round(ci_lo, 4), round(ci_hi, 4)],
        "kappa_interpretation": interp,
    }
    json_path = OUTPUT_DIR / f"resumo_{ts}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(f"  Resumo JSON: {json_path.name}\n")


if __name__ == "__main__":
    main()
